from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from PIL import Image
from docx import Document as DocxDocument
from pypdf import PdfReader
from pytesseract import pytesseract, TesseractNotFoundError
from ..config import get_settings
from ..models import Document
from .embedding_service import EmbeddingService
from .supabase_service import SupabaseService


class DocumentService:
    def __init__(
        self,
        supabase: SupabaseService | None = None,
        embedding_service: EmbeddingService | None = None,
    ):
        self.supabase = supabase or SupabaseService()
        self.settings = get_settings()
        self.embedding_service = embedding_service or EmbeddingService(self.supabase)

    def list_documents(self, case_id: str) -> list[Document]:
        records = self.supabase.select(
            "documents", filters={"case_id": case_id}, order=("created_at", False)
        )
        return [Document(**record) for record in records]

    def upload_file(self, bucket: str, file_path: str, file_bytes: bytes, content_type: str) -> str:
        object_path = f"{uuid4()}/{file_path}"
        storage_bucket = self.supabase.client.storage.from_(bucket)
        storage_bucket.upload(object_path, file_bytes, {"content-type": content_type})
        public_url = storage_bucket.get_public_url(object_path)
        return public_url["publicUrl"]

    def delete_document(self, document_id: str) -> None:
        self.supabase.client.table("document_embeddings").delete().eq("document_id", document_id).execute()
        self.supabase.delete("documents", document_id)

    def extract_text_from_upload(self, file_name: str, content_type: str | None, file_bytes: bytes) -> str | None:
        suffix = Path(file_name).suffix.lower()
        mime = content_type or ""
        try:
            if suffix == ".txt" or mime.startswith("text/"):
                return file_bytes.decode("utf-8", errors="ignore")
            if suffix == ".pdf" or mime in {"application/pdf"}:
                reader = PdfReader(BytesIO(file_bytes))
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            if suffix in {".docx"} or mime in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}:
                doc = DocxDocument(BytesIO(file_bytes))
                return "\n".join(paragraph.text for paragraph in doc.paragraphs)
            if suffix in {".doc"}:
                # Legacy doc not natively supported; attempt text decode
                return file_bytes.decode("latin-1", errors="ignore")
            if mime.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
                try:
                    image = Image.open(BytesIO(file_bytes))
                    return pytesseract.image_to_string(image, lang="rus+eng")
                except TesseractNotFoundError:
                    logging.warning("Tesseract OCR engine not found; skipping OCR for %s", file_name)
                    return None
                except Exception as exc:  # pragma: no cover - Pillow-specific errors
                    logging.warning("Failed to process image %s: %s", file_name, exc)
                    return None
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception as exc:  # pragma: no cover - extraction failure
            logging.warning("Failed to extract text from %s: %s", file_name, exc)
            return None

    def save_document_metadata(
        self,
        case_id: str,
        file_name: str,
        file_url: str,
        content: str | None,
    ) -> Document:
        record = self.supabase.insert(
            "documents",
            {
                "case_id": case_id,
                "file_name": file_name,
                "file_url": file_url,
                "content": content,
            },
        )[0]
        return Document(**record)

    def store_document(
        self,
        *,
        case_id: str,
        file_name: str,
        content_type: str | None,
        file_bytes: bytes,
    ) -> Document:
        text_content = self.extract_text_from_upload(file_name, content_type, file_bytes)
        file_url = self.upload_file(
            "legal-assistant-uploads",
            file_name,
            file_bytes,
            content_type or "application/octet-stream",
        )
        document = self.save_document_metadata(case_id, file_name, file_url, text_content)
        if text_content:
            self.embedding_service.reindex_document(document, text_content)
        return document

